#python to docx
from docx import Document

def Wiki(keyword, lang = 'th'):
    wikipedia.set_lang(lang)
    # summary
    data = wikipedia.summary(keyword)

    #page + content
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() #สร้างไฟล์ word ใน python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('สร้างไฟล์สำเร็จ')


import wikipedia
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
    
GUI = Tk()
GUI.title('GUI_Wiki')
GUI.geometry('500x500')

#config
FONT = ('Angsana New',15)

#label
L = ttk.Label(GUI, text = 'Input word to search', font = FONT)
L.pack()

#Search input field
v_search = StringVar()
E = ttk.Entry(GUI, textvariable = v_search, font = FONT)
E.pack(pady=10)


#Search button
def Search():
    keyword = v_search.get()
    try:
        language = v_radio.get()
        Wiki(keyword, language)
        messagebox.showinfo('saved','File is created successfully')
    except:
        messagebox.showwarning('Keyword Error', 'Keyword not found, try again')


B = ttk.Button(GUI, text = 'Search',command = Search)
B.pack(ipadx=20, ipady=10)

# เลือกภาษา
F = Frame(GUI)
F.pack(pady=10)

v_radio = StringVar()

RB1 = ttk.Radiobutton(F,text='Thai',variable = v_radio,value = 'th')
RB2 = ttk.Radiobutton(F,text='English',variable = v_radio,value = 'en')
RB3 = ttk.Radiobutton(F,text='Chinese',variable = v_radio,value = 'zh')
RB1.invoke() #สั่งให้ค่าเริ่มต้นเป็นภาษาไทย

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)



GUI.mainloop()
